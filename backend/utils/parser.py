import io
import re
import fitz  # PyMuPDF
import docx  # python-docx
from typing import Dict, Any, List, Optional, Set

# Master Technical & Domain Skills Vocabulary Dictionary
MASTER_SKILLS_DICTIONARY: List[str] = [
    "react", "typescript", "javascript", "python", "fastapi", "django", "flask",
    "node.js", "express", "java", "spring boot", "c++", "c#", ".net", "php", "sql",
    "postgresql", "mysql", "mongodb", "redis", "docker", "kubernetes", "aws", "azure",
    "gcp", "git", "github", "ci/cd", "rest apis", "graphql", "tailwind css", "html",
    "css", "next.js", "redux", "pytorch", "tensorflow", "scikit-learn", "pandas",
    "numpy", "microservices", "system design", "linux", "unit testing", "scrum",
    "agile", "jira", "leadership", "communication", "problem solving", "machine learning",
    "deep learning", "nlp", "computer vision", "opencv", "kafka", "elasticsearch", "tableau",
    "power bi", "excel", "spark", "hadoop", "bash", "shell", "terraform", "ansible"
]

LANGUAGES_DICTIONARY: List[str] = [
    "english", "spanish", "french", "german", "mandarin", "chinese", "hindi",
    "japanese", "italian", "portuguese", "russian", "arabic"
]

MAX_FILE_SIZE_BYTES = 10 * 1024 * 1024  # 10 MB

NON_NAME_HEADER_WORDS = {
    "resume", "curriculum", "vitae", "cv", "summary", "experience", "education",
    "skills", "projects", "certifications", "contact", "profile", "objective",
    "engineer", "developer", "manager", "page", "phone", "email", "address",
    "technical", "technologies", "competencies", "work", "history"
}

SECTION_HEADER_PATTERNS = {
    "skills": r'\b(?:technical\s+skills|skills|technologies|core\s+competencies|tools\s+&\s+technologies|programming\s+languages)\b',
    "education": r'\b(?:education|academic\s+background|academic\s+qualifications|educational\s+history)\b',
    "experience": r'\b(?:experience|work\s+experience|employment\s+history|professional\s+experience|work\s+history)\b',
    "projects": r'\b(?:projects|key\s+projects|personal\s+projects|academic\s+projects)\b',
    "certifications": r'\b(?:certifications|certificates|licenses|professional\s+certifications)\b',
    "languages": r'\b(?:languages|spoken\s+languages|language\s+proficiency)\b'
}

def clean_text(raw_text: str) -> str:
    if not raw_text:
        return ""
    lines = [line.strip() for line in raw_text.splitlines() if line.strip()]
    text = " ".join(lines)
    return re.sub(r'\s+', ' ', text).strip()

def extract_sections(text: str) -> Dict[str, str]:
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    sections: Dict[str, List[str]] = {
        "skills": [],
        "education": [],
        "experience": [],
        "projects": [],
        "certifications": [],
        "languages": [],
        "other": []
    }

    current_section = "other"

    for line in lines:
        line_lower = line.lower()
        matched_new_section = False

        if len(line.split()) <= 4:
            for sec_name, pattern in SECTION_HEADER_PATTERNS.items():
                if re.search(pattern, line_lower):
                    current_section = sec_name
                    matched_new_section = True
                    break

        if not matched_new_section:
            sections[current_section].append(line)

    return {sec: " ".join(content) for sec, content in sections.items()}

def extract_name(text: str, filename: str = "") -> Optional[str]:
    if not text:
        return None

    lines = [line.strip() for line in text.splitlines() if line.strip()]

    if len(lines) <= 2:
        candidates = text.split(" ")
        possible_lines = [
            " ".join(candidates[:2]),
            " ".join(candidates[:3]),
            " ".join(candidates[:4]),
            " ".join(candidates[1:3]),
            " ".join(candidates[1:4])
        ]
        lines.extend(possible_lines)

    for line in lines[:8]:
        line_clean = line.strip()

        if re.search(r'[@\d:]|http|www|\.com|\.io|\.org', line_clean, re.IGNORECASE):
            continue

        words = line_clean.split()
        if not (2 <= len(words) <= 4):
            continue

        lower_words = [w.lower().strip(".,-") for w in words]
        if any(w in NON_NAME_HEADER_WORDS for w in lower_words):
            continue

        if all(re.match(r'^[A-Za-z\.\'-]+$', w) for w in words):
            name_candidate = " ".join(words).title()
            if 3 <= len(name_candidate) <= 40:
                return name_candidate

    return None

def extract_email(text: str) -> Optional[str]:
    pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    match = re.search(pattern, text)
    if match:
        return match.group(0).rstrip(".,;)]>")
    return None

def extract_phone(text: str) -> Optional[str]:
    patterns = [
        r'\+?\d{1,3}[-.\s]?\(?\d{2,4}\)?[-.\s]?\d{3,4}[-.\s]?\d{3,4}',
        r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
        r'\+\d{1,3}\s?\d{4,5}\s?\d{4,5}'
    ]
    for p in patterns:
        matches = re.finditer(p, text)
        for match in matches:
            ph = match.group(0).strip()
            digits = re.sub(r'\D', '', ph)
            if 7 <= len(digits) <= 15:
                return ph
    return None

def extract_linkedin(text: str) -> Optional[str]:
    pattern = r'(?:https?:\/\/)?(?:www\.)?linkedin\.com\/in\/[a-zA-Z0-9_-]+\/?'
    match = re.search(pattern, text, re.IGNORECASE)
    if match:
        url = match.group(0).rstrip("/")
        if not url.startswith("http"):
            url = "https://" + url
        return url
    return None

def extract_github(text: str) -> Optional[str]:
    pattern = r'(?:https?:\/\/)?(?:www\.)?github\.com\/[a-zA-Z0-9_-]+\/?'
    matches = re.finditer(pattern, text, re.IGNORECASE)
    ignored_slugs = {"features", "pricing", "about", "topics", "trending", "explore", "enterprise"}

    for match in matches:
        url = match.group(0).rstrip("/")
        slug = url.split("/")[-1].lower()
        if slug not in ignored_slugs and len(slug) > 1:
            if not url.startswith("http"):
                url = "https://" + url
            return url
    return None

def extract_skills(text: str, sections: Dict[str, str]) -> List[str]:
    found_skills: Set[str] = set()

    text_lower = text.lower()
    for sk in MASTER_SKILLS_DICTIONARY:
        if re.search(r'\b' + re.escape(sk) + r'\b', text_lower):
            found_skills.add(sk.title())

    skills_sec = sections.get("skills", "")
    if skills_sec:
        items = re.split(r'[,|;•·\n\t]', skills_sec)
        for item in items:
            cleaned_item = item.strip().title()
            if 2 <= len(cleaned_item) <= 30 and not re.search(r'[@\d]', cleaned_item):
                if cleaned_item.lower() not in NON_NAME_HEADER_WORDS:
                    found_skills.add(cleaned_item)

    return sorted(list(found_skills))

def extract_education(text: str, sections: Dict[str, str]) -> List[str]:
    edu_text = sections.get("education", "") or text
    pattern = r'\b(?:bachelor|master|b\.s\.|m\.s\.|ph\.d\.|b\.tech|m\.tech|degree|diploma|associate|university|college|institute)\b[^\.\n]*'
    matches = re.findall(pattern, edu_text, re.IGNORECASE)
    cleaned = [m.strip().title() for m in matches if len(m.strip()) > 5]
    return list(dict.fromkeys(cleaned))

def extract_experience(text: str, sections: Dict[str, str]) -> List[str]:
    exp_text = sections.get("experience", "") or text
    pattern = r'\b(?:senior|junior|lead|developer|engineer|manager|consultant|analyst|intern|architect|specialist)\b[^\.\n]*'
    matches = re.findall(pattern, exp_text, re.IGNORECASE)
    cleaned = [m.strip().title() for m in matches if len(m.strip()) > 8]
    return list(dict.fromkeys(cleaned))[:4]

def extract_projects(text: str, sections: Dict[str, str]) -> List[str]:
    proj_text = sections.get("projects", "") or text
    pattern = r'\b(?:project|built|developed|architected|created|designed|implemented)\b[^\.\n]*'
    matches = re.findall(pattern, proj_text, re.IGNORECASE)
    cleaned = [m.strip().title() for m in matches if len(m.strip()) > 10]
    return list(dict.fromkeys(cleaned))[:4]

def extract_certifications(text: str, sections: Dict[str, str]) -> List[str]:
    cert_text = sections.get("certifications", "") or text
    pattern = r'\b(?:certified|certification|certificate|aws certified|coursera|udemy|scrum master|pmp|cka|cissp)\b[^\.\n]*'
    matches = re.findall(pattern, cert_text, re.IGNORECASE)
    cleaned = [m.strip().title() for m in matches if len(m.strip()) > 5]
    return list(dict.fromkeys(cleaned))[:4]

def extract_languages(text: str, sections: Dict[str, str]) -> List[str]:
    lang_text = sections.get("languages", "") or text
    text_lower = lang_text.lower()
    matched = []
    for lang in LANGUAGES_DICTIONARY:
        if re.search(r'\b' + re.escape(lang) + r'\b', text_lower):
            matched.append(lang.title())
    return sorted(list(set(matched)))

def parse_resume_structured(text: str, filename: str = "") -> Dict[str, Any]:
    sections = extract_sections(text)

    return {
        "name": extract_name(text, filename),
        "email": extract_email(text),
        "phone": extract_phone(text),
        "linkedin": extract_linkedin(text),
        "github": extract_github(text),
        "skills": extract_skills(text, sections),
        "education": extract_education(text, sections),
        "experience": extract_experience(text, sections),
        "projects": extract_projects(text, sections),
        "certifications": extract_certifications(text, sections),
        "languages": extract_languages(text, sections)
    }

def extract_text_from_pdf(pdf_source) -> str:
    raw_text = ""
    try:
        if isinstance(pdf_source, bytes):
            doc = fitz.open(stream=pdf_source, filetype="pdf")
        elif isinstance(pdf_source, (str, io.BytesIO)):
            doc = fitz.open(pdf_source)
        else:
            doc = fitz.open(stream=pdf_source.read(), filetype="pdf")

        for page in doc:
            raw_text += page.get_text() + "\n"
        doc.close()
    except fitz.FileDataError:
        raise ValueError("The uploaded file is an invalid or corrupted PDF document. Please upload a clean PDF file.")
    except Exception as e:
        raise ValueError(f"Failed to parse PDF document: {str(e)}")

    return clean_text(raw_text)

def extract_text_from_docx(docx_source) -> str:
    raw_text = ""
    try:
        if isinstance(docx_source, bytes):
            file_stream = io.BytesIO(docx_source)
            doc = docx.Document(file_stream)
        else:
            doc = docx.Document(docx_source)

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        raw_text = "\n".join(paragraphs)
    except Exception as e:
        raise ValueError(f"The uploaded file is an invalid or corrupted DOCX document: {str(e)}")

    return clean_text(raw_text)

def validate_and_parse_file(file_bytes: bytes, filename: str) -> str:
    """
    Robust File Validator & Text Extractor:
    Handles:
    - Empty uploads (0-byte file)
    - Large files (>10MB)
    - Unsupported file formats (.png, .zip, .exe)
    - Invalid & Corrupted PDFs
    - Corrupted DOCX files
    - Blank / image-only PDFs with no text
    """
    if not file_bytes or len(file_bytes) == 0:
        raise ValueError("The uploaded file is empty (0 bytes). Please upload a valid PDF or DOCX resume document.")

    if len(file_bytes) > MAX_FILE_SIZE_BYTES:
        mb_size = len(file_bytes) / (1024 * 1024)
        raise ValueError(f"File size exceeds maximum 10MB limit ({mb_size:.1f} MB uploaded). Please upload a smaller file.")

    ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ""
    if ext not in ["pdf", "docx", "doc"]:
        raise ValueError(f"Unsupported file format '.{ext}'. Only PDF and DOCX resume documents are supported.")

    if ext == "pdf":
        text = extract_text_from_pdf(file_bytes)
    else:
        text = extract_text_from_docx(file_bytes)

    if not text or len(text.strip()) < 10:
        raise ValueError("No readable text could be extracted from the resume. Please ensure the file contains selectable text and is not blank or image-only.")

    return text

def parse_resume_file(file_source, filename: str) -> str:
    if isinstance(file_source, bytes):
        return validate_and_parse_file(file_source, filename)
    filename_lower = filename.lower()
    if filename_lower.endswith(".pdf"):
        return extract_text_from_pdf(file_source)
    elif filename_lower.endswith(".docx") or filename_lower.endswith(".doc"):
        return extract_text_from_docx(file_source)
    else:
        if isinstance(file_source, bytes):
            return clean_text(file_source.decode('utf-8', errors='ignore'))
        return ""
